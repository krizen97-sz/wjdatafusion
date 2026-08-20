#!/usr/bin/env python3
"""Build the deployment and operation manuals for the rynew delivery bundle."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK = "17365D"
LIGHT = "EAF2F8"
MUTED = "667085"
GREEN = "2E7D32"
RED = "B42318"
PAGE_WIDTH_TWIPS = 9360
CELL_MARGIN_TWIPS = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=CELL_MARGIN_TWIPS, start=CELL_MARGIN_TWIPS,
                     bottom=CELL_MARGIN_TWIPS, end=CELL_MARGIN_TWIPS) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def configure_table_geometry(table, widths: list[int], indent_twips: int = CELL_MARGIN_TWIPS) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col in list(grid):
        grid.remove(col)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    paragraph.add_run(" 页").font.size = Pt(9)


def add_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids or [0]) + 1
    next_num = max(num_ids or [0]) + 1

    def create(abstract_id: int, num_id: int, fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "420")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "420")
        ind.set(qn("w:hanging"), "240")
        p_pr.append(ind)
        lvl.append(p_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create(next_abstract, next_num, "decimal", "%1.")
    create(next_abstract + 1, next_num + 1, "bullet", "•")
    return next_num, next_num + 1


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def restart_numbering(document: Document, source_num_id: int) -> int:
    """Create a fresh numbering instance so each numbered procedure starts at 1."""
    numbering = document.part.numbering_part.element
    source = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == source_num_id
    )
    abstract_id = source.find(qn("w:abstractNumId")).get(qn("w:val"))
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_num_id = max(num_ids or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return new_num_id


class Manual:
    def __init__(self, title: str, subtitle: str, version: str):
        self.doc = Document()
        self.num_id, self.bullet_id = add_numbering(self.doc)
        self._configure_styles()
        self._configure_section()
        self._cover(title, subtitle, version)

    def _configure_styles(self) -> None:
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Arial Unicode MS"
        normal.font.size = Pt(11)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.25
        for name, size, before, after in (("Title", 28, 0, 18), ("Subtitle", 14, 0, 12),
                                          ("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7),
                                          ("Heading 3", 12, 10, 5)):
            style = styles[name]
            style.font.name = "Arial Unicode MS"
            style.font.size = Pt(size)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
            style.font.color.rgb = RGBColor.from_string(DARK if name in ("Title", "Heading 1") else BLUE)
            style.font.bold = name != "Subtitle"
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
        styles["Caption"].font.name = "Arial Unicode MS"
        styles["Caption"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        styles["Caption"].font.size = Pt(9)
        styles["Caption"].font.color.rgb = RGBColor.from_string(MUTED)

    def _configure_section(self) -> None:
        sec = self.doc.sections[0]
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.35)
        sec.footer_distance = Inches(0.35)
        header = sec.header.paragraphs[0]
        header.text = "华东信息融合平台 · 交付文档"
        header.style = self.doc.styles["Caption"]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_field(sec.footer.paragraphs[0])

    def _cover(self, title: str, subtitle: str, version: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("RYNEW")
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor.from_string(BLUE)
        title_p = self.doc.add_paragraph(title, style="Title")
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p = self.doc.add_paragraph(subtitle, style="Subtitle")
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rule = self.doc.add_paragraph("━━━━━━━━━━━━━━━━━━━━")
        rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rule.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
        for text in (f"交付版本：{version}", "编制日期：2026 年 8 月 10 日", "适用范围：前端、后端、数据库与业务模块"):
            q = self.doc.add_paragraph(text)
            q.alignment = WD_ALIGN_PARAGRAPH.CENTER
            q.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        self.doc.add_paragraph().paragraph_format.space_before = Pt(90)
        note = self.doc.add_paragraph("内部交付资料｜部署前请完成环境参数与密钥替换")
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        self.doc.add_page_break()

    def h1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)

    def p(self, text: str, bold_prefix: str | None = None) -> None:
        p = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            p.add_run(bold_prefix).bold = True
            p.add_run(text[len(bold_prefix):])
        else:
            p.add_run(text)

    def list(self, items: list[str], numbered: bool = False) -> None:
        list_num_id = restart_numbering(self.doc, self.num_id) if numbered else self.bullet_id
        for item in items:
            p = self.doc.add_paragraph()
            apply_num(p, list_num_id)
            p.add_run(item)

    def callout(self, title: str, text: str, tone: str = "blue") -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        configure_table_geometry(table, [PAGE_WIDTH_TWIPS], 180)
        cell = table.cell(0, 0)
        set_cell_width(cell, PAGE_WIDTH_TWIPS)
        set_cell_shading(cell, {"blue": LIGHT, "green": "EAF7EE", "red": "FDECEC"}[tone])
        set_cell_margins(cell, 160, 180, 160, 180)
        p = cell.paragraphs[0]
        r = p.add_run(f"{title}  ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string({"blue": BLUE, "green": GREEN, "red": RED}[tone])
        p.add_run(text)
        set_repeat_table_header(table.rows[0])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def table(self, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        widths = widths or [PAGE_WIDTH_TWIPS // len(headers)] * len(headers)
        configure_table_geometry(table, widths)
        for i, (head, width) in enumerate(zip(headers, widths)):
            cell = table.rows[0].cells[i]
            set_cell_width(cell, width)
            set_cell_shading(cell, DARK)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cell.paragraphs[0].add_run(head)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        set_repeat_table_header(table.rows[0])
        prevent_row_split(table.rows[0])
        for row_idx, row in enumerate(rows):
            cells = table.add_row().cells
            for i, (value, width) in enumerate(zip(row, widths)):
                set_cell_width(cells[i], width)
                set_cell_margins(cells[i])
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_idx % 2:
                    set_cell_shading(cells[i], "F7F9FC")
                cells[i].paragraphs[0].add_run(str(value))
            prevent_row_split(table.rows[-1])
        self.doc.add_paragraph()

    def image(self, path: Path, caption: str, width: float = 6.45) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        # Keep business screenshots focused on the actual operation area. The
        # dynamic navigation tree belongs to the source environment and may
        # contain modules that are intentionally absent from this delivery.
        crop_left = 12000 if path.name.startswith(("site-", "auto-")) else 0
        crop_bottom = 9000 if path.name.startswith("00-login") else 0
        if crop_left or crop_bottom:
            blip_fill = run._r.xpath(".//pic:blipFill")[0]
            src_rect = OxmlElement("a:srcRect")
            if crop_left:
                src_rect.set("l", str(crop_left))
            if crop_bottom:
                src_rect.set("b", str(crop_bottom))
            blip_fill.insert(1, src_rect)
        descr = OxmlElement("wp:docPr")
        # python-docx already creates docPr; set its description for accessibility.
        for node in run._r.xpath(".//wp:docPr"):
            node.set("descr", caption)
        c = self.doc.add_paragraph(caption, style="Caption")
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def page_break(self) -> None:
        self.doc.add_page_break()

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        # LibreOffice headless may ignore style-level East Asian font hints.
        # Apply a CJK-capable font directly to every run for deterministic PDF rendering.
        containers = [self.doc]
        for section in self.doc.sections:
            containers.extend([section.header, section.footer])
        for container in containers:
            for paragraph in container.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Hiragino Sans GB"
                    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                        fonts.set(qn(f"w:{attr}"), "Hiragino Sans GB")
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Hiragino Sans GB"
                                fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
                                for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                                    fonts.set(qn(f"w:{attr}"), "Hiragino Sans GB")
        props = self.doc.core_properties
        props.title = self.doc.paragraphs[1].text if len(self.doc.paragraphs) > 1 else "Rynew Manual"
        props.subject = "华东信息融合平台交付资料"
        props.author = "Rynew 交付团队"
        props.keywords = "Rynew, 部署, 操作手册"
        props.comments = "Sanitized delivery document; no passwords are embedded."
        self.doc.save(path)


def build_deployment(out: Path, shots: Path) -> None:
    m = Manual("前后端程序部署手册", "后端脚本 · 外部配置 · 前端目录", "v3.9.1 · 20260810 现场融合/自动巡检交付包")
    m.h1("1. 文档目的与交付边界")
    m.p("本手册只保留现场部署最常用的步骤：导入数据库、填写两个 Spring 配置文件、运行一个 Java 脚本、解压前端 dist 并放置 Nginx 配置文件。交付范围为平台基础业务、现场融合管理和现场自动巡检。")
    m.callout("最简路径", "数据库初始化 → 修改 config/external 下的配置 → 执行 start-backend.sh → 解压前端 dist → 放置 Nginx 配置并 reload。")
    m.table(["交付物", "用途", "默认目标位置"], [
        ["frontend/rynew-frontend-v3.9.1-20260810.tar.gz", "Vue 静态站点", "/opt/rynew/frontend"],
        ["backend/wjdatafusion-admin-v3.9.1-20260810.jar", "Spring Boot 后端", "/opt/rynew/backend"],
        ["database/rynew-init-v3.9.1-site-auto-demo.sql", "表结构、基础配置与虚构演示数据", "MySQL rynew 库"],
        ["config/external/application*.yml", "现场外部配置模板", "/opt/rynew/config"],
        ["config/start-backend.sh", "后端直接运行脚本", "/opt/rynew/start-backend.sh"],
        ["config/rynew-nginx.conf", "前端与接口代理示例", "/etc/nginx/conf.d/rynew.conf"],
    ], [2500, 3200, 3660])

    m.h1("2. 文件准备")
    m.list([
        "校验 SHA256SUMS.txt，确认文件传输完整。",
        "创建 /opt/rynew/backend、/opt/rynew/frontend、/opt/rynew/config、/opt/rynew/upload、/opt/rynew/logs。",
        "复制 JAR、两个 YAML、start-backend.sh 和前端压缩包到对应目录。",
    ], numbered=True)
    m.callout("目录命令", "install -d /opt/rynew/{backend,frontend,config,upload,logs}")

    m.image(shots / "deploy-01-directory.png", "图 1  交付文件与现场目标目录对应关系")

    m.h1("3. 数据库初始化")
    m.h2("3.1 数据范围")
    m.p("初始化 SQL 包含平台当前有效表结构、平台运行所需基础字典/菜单/角色配置，以及现场融合管理、现场自动巡检的虚构演示数据。SQL 不包含线上真实现场、联系人、设备地址、登录凭据、巡检目标或历史执行记录。")
    m.table(["模块", "演示数据特征", "明确不包含"], [
        ["现场融合", "演示现场、10.255.253.x、13800000000", "真实地址、联系人、凭据"],
        ["自动巡检", "example.invalid 目标、停用计划", "真实设备地址、用户名、密码、执行记录"],
    ], [1800, 3400, 4160])
    m.h2("3.2 初始化步骤")
    m.list([
        "创建空库：CREATE DATABASE rynew CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;",
        "导入：mysql --default-character-set=utf8mb4 -u root -p rynew < rynew-init-v3.9.1-site-auto-demo.sql",
        "执行脚本末尾校验查询，确认核心表、演示数据和禁用巡检计划均存在。",
        "首次登录后立即修改管理员密码，并按岗位重新分配权限。",
    ], numbered=True)
    m.callout("注意", "本 SQL 面向空库初始化。已有生产库必须先做版本差异评审，不可直接覆盖导入。", "red")

    m.h1("4. 后端外部配置")
    m.h2("4.1 application.yml：端口、Redis 与密钥")
    m.p("复制交付包 config/external/application.yml 到 /opt/rynew/config/application.yml，然后只修改现场值。配置文件位于 JAR 外部，升级程序时无需重新打包。")
    m.image(shots / "deploy-02-application-yml.png", "图 2  application.yml 实际文件：重点填写端口、目录、Redis 和两项密钥")
    m.table(["配置项", "现场填写", "说明"], [
        ["server.port", "8080 或现场端口", "须与 Nginx proxy_pass 一致"],
        ["ruoyi.profile / log.path", "上传目录 / 日志目录", "目录需提前创建并可写"],
        ["spring.data.redis.host / port", "Redis 内网 IP / 端口", "从应用服务器测试可达"],
        ["database / password", "现场分配索引 / 密码", "多节点必须一致"],
        ["token.secret", "随机 64 字符以上", "所有后端节点必须一致"],
        ["support.credential.key", "现场独立 AES 密钥", "上线后不可随意更换"],
    ], [3100, 3100, 3160])
    m.callout("必须替换", "TOKEN_SECRET 和 SUPPORT_CREDENTIAL_KEY 不得保留 CHANGE_ME；口令和密钥不要写进聊天、工单截图或版本库。", "red")

    m.h2("4.2 application-druid.yml：MySQL")
    m.image(shots / "deploy-03-application-druid.png", "图 3  application-druid.yml 实际文件：填写数据库地址、库名、账号和密码")
    m.table(["配置项", "现场填写", "说明"], [
        ["spring.datasource.druid.master.url", "jdbc:mysql://数据库IP:3306/rynew?...", "库名、时区、字符集按现场确认"],
        ["spring.datasource.druid.master.username", "专用应用账号", "只授予 rynew 库必要权限"],
        ["spring.datasource.druid.master.password", "现场应用账号密码", "配置文件权限建议 600"],
        ["maxActive / minIdle", "20 / 5 起步", "按数据库连接上限和应用并发调整"],
    ], [3100, 3300, 2960])
    m.p("推荐 URL：jdbc:mysql://数据库IP:3306/rynew?useUnicode=true&characterEncoding=utf8&zeroDateTimeBehavior=convertToNull&useSSL=false&serverTimezone=GMT%2B8")

    m.h1("5. 后端启动脚本")
    m.p("脚本直接执行 java -jar，并在 JAR 后追加 Spring Boot 参数，指定外部配置目录。additional-location 会在保留包内默认配置的同时加载外部文件。")
    m.image(shots / "deploy-04-start-script.png", "图 4  start-backend.sh 实际文件：java -jar 与外部 Spring 配置路径")
    m.callout("关键命令", "java -jar /opt/rynew/backend/wjdatafusion-admin-v3.9.1-20260810.jar --spring.config.additional-location=optional:file:/opt/rynew/config/")
    m.table(["参数", "作用", "注意"], [
        ["-jar ...jar", "直接运行后端程序", "路径必须指向当前交付 JAR"],
        ["--spring.config.additional-location", "加载 JAR 外部配置", "目录末尾必须保留 /"],
        ["optional:file:", "配置缺失时给出正常 Spring 处理", "生产仍应确认两个 YAML 均存在"],
    ], [3100, 3100, 3160])
    m.list([
        "授权并启动：chmod 750 /opt/rynew/start-backend.sh && /opt/rynew/start-backend.sh",
        "查看日志：tail -f /opt/rynew/logs/backend-console.log",
        "确认启动：日志出现 Started RuoYiApplication，且无 MySQL/Redis 连接错误。",
    ], numbered=True)
    m.h2("5.1 停止与检查")
    m.list([
        "查看进程：ps -ef | grep '[w]jdatafusion-admin'",
        "查看日志：tail -f /opt/rynew/logs/backend-console.log",
        "优雅停止：前台运行按 Ctrl+C；后台运行记录 PID 后执行 kill <PID>。",
        "端口检查：ss -lntp | grep 8080；接口检查：curl -I http://127.0.0.1:8080/captchaImage。",
    ])
    m.callout("不要使用", "不要用 kill -9 作为日常停止命令；只有普通 kill 长时间无效并确认无任务执行时才升级处理。", "red")

    m.h1("6. 前端 dist 与 Nginx 配置")
    m.p("本手册不包含 Nginx 安装。只要现场已有 Nginx，将 dist 解压到目录，并放置交付的配置文件即可。")
    m.image(shots / "deploy-05-frontend-dist.png", "图 5  前端 dist 解压后的目录：index.html 必须位于根目录")
    m.list([
        "解压前端包到空的 /opt/rynew/frontend 目录，确保 index.html 位于目录根部。",
        "复制 config/rynew-nginx.conf 到现场 Nginx 站点配置目录，修改 root 和 proxy_pass。",
        "执行 nginx -t，确认语法通过后 reload；不要先覆盖正在使用的站点配置。",
        "浏览器访问站点，确认登录页、静态资源和 /prod-api/captchaImage 均正常。",
    ], numbered=True)
    m.image(shots / "deploy-06-nginx.png", "图 6  rynew-nginx.conf 实际文件：root 指向 dist，proxy_pass 指向后端")
    m.p("/prod-api/ 的 proxy_pass 末尾保留 /，用于移除前缀并转发至后端根路径。修改后执行 nginx -t，成功后再 nginx -s reload。")

    m.h1("7. 验证与回滚")
    m.list(["MySQL、Redis", "后端 JAR", "Nginx reload", "登录与业务模块验证"], numbered=True)
    m.table(["检查项", "方法", "通过标准"], [
        ["进程", "ps -ef | grep '[w]jdatafusion-admin'", "后端 Java 进程存在"],
        ["端口", "ss -lntp | grep -E ':80|:8080'", "80 对外、8080 仅内网/本机"],
        ["登录", "访问登录页并获取验证码", "无 404/502，能正常登录"],
        ["基础业务", "用户、角色、字典只读查看", "列表加载与权限符合预期"],
        ["现场融合", "打开现场列表与配置画布", "演示现场可查看，不出现真实凭据"],
        ["自动巡检", "查看模板、目标、计划", "演示计划为停用；人工启用后才可执行"],
    ], [1800, 3600, 3960])
    m.callout("验收证据", "记录部署时间、制品 SHA-256、数据库备份号、验证账号、每个模块的页面截图和操作结果。")

    m.h2("7.1 简单回滚")
    m.list([
        "停止后端进程，确认没有正在执行的自动巡检任务。",
        "恢复后端：把 JAR 路径切回上一版本，并使用同一外部配置启动。",
        "恢复前端：将 Nginx root 切回上一份带时间戳的静态目录并 reload。",
        "恢复数据库：仅在已确认数据库变更不向后兼容时，从上线前备份恢复；先在隔离环境验证恢复点。",
        "复核：重新执行登录、基础业务、现场融合和自动巡检的最小回归，并保存日志。",
    ], numbered=True)
    m.callout("不可逆风险", "数据库恢复会覆盖上线后的新增业务数据，必须经业务负责人确认恢复时间点和影响范围。", "red")

    m.h2("7.2 常见问题")
    m.table(["现象", "排查顺序", "处理建议"], [
        ["Nginx 502", "后端进程 → 8080 → proxy_pass", "先恢复后端健康，再 reload Nginx"],
        ["登录后立即退出", "Token secret → Redis → 系统时间", "保持多节点密钥一致并校时"],
        ["数据库连接失败", "URL/账号 → 防火墙 → TLS/时区", "用应用账户最小化测试 SELECT 1"],
        ["现场/巡检凭据无法解密", "密钥变量 → 历史密钥 → 数据来源", "不要直接换密钥；先备份并制定轮换"],
        ["巡检不执行", "计划状态 → Quartz → 目标/模板 → 日志", "先用停用演示计划做手动测试"],
        ["页面刷新 404", "Nginx try_files", "确保回退到 /index.html"],
    ], [2100, 3300, 3960])

    m.save(out)


def build_operation(out: Path, shots: Path) -> None:
    m = Manual("平台操作手册", "基础业务 · 现场融合管理 · 现场自动巡检", "v3.9.1 · 20260810 现场融合/自动巡检交付包")
    m.h1("1. 使用说明")
    m.p("本手册面向平台管理员、现场运维人员和巡检人员。页面截图来自当前登录页及既有真实操作过程，用于说明真实页面路径与按钮位置；初始化 SQL 中的业务数据已替换为虚构演示数据。由于权限和部署版本不同，菜单名称与按钮可见性可能略有差异。")
    m.callout("安全原则", "涉及删除、批量修改、凭据查看、巡检启停等操作时，先确认目标范围和权限，再执行并保留操作记录。", "red")
    m.table(["角色", "主要职责", "建议权限"], [
        ["系统管理员", "用户、角色、菜单、参数、日志", "系统管理；不默认授予业务凭据查看"],
        ["现场运维", "现场、平台、组织、设备位置", "support:site/server/platform 等"],
        ["巡检管理员", "模板、工具、目标、计划", "support:autoInspection:* 与计划权限"],
        ["审计/只读", "查看报表、日志、结果", "仅 list/query/export，禁止变更"],
    ], [1800, 3500, 4060])
    m.image(shots / "00-login-current.png", "图 1  当前构建登录页：输入账号、密码和验证码后进入平台")

    m.h1("2. 平台基础业务操作")
    m.h2("2.1 登录与导航")
    m.list([
        "使用管理员分配的账号登录；连续输错密码可能触发锁定。",
        "通过左侧菜单进入业务模块；顶部标签用于在多个页面间切换。",
        "首次登录或管理员重置后立即修改密码，不在浏览器保存共享账号。",
        "退出时使用右上角退出功能；公共终端关闭浏览器并清除缓存。",
    ], numbered=True)
    m.h2("2.2 用户、角色与权限")
    m.list([
        "系统管理 → 用户管理：新增用户，绑定部门、岗位和角色；停用优先于删除。",
        "系统管理 → 角色管理：按岗位勾选菜单和按钮权限，避免直接复用超级管理员。",
        "系统管理 → 菜单管理：检查新增业务菜单的路径、组件和权限字符。",
        "系统监控 → 操作日志/登录日志：按用户、时间、状态筛选并导出审计证据。",
    ], numbered=True)
    m.callout("权限生效", "角色或菜单调整后，用户应退出并重新登录；若仍不生效，再检查 Redis 会话与前端缓存。")
    m.h2("2.3 通用列表与表单")
    m.p("列表页通常提供查询、重置、新增、修改、删除、导入和导出。带红色星号为必填项。导入前先下载模板，小批量试导后再执行正式导入；删除前确认关联关系和备份。")

    m.h1("3. 现场融合管理模块")
    m.p("现场融合管理以“现场”为主线，把现场基础资料、平台、组织、联系人、服务器、设备机房/机柜、硬件资产和版本记录集中维护。建议先建现场，再逐步补齐节点并核对关系。")
    m.h2("3.1 现场列表、查询与状态管理")
    m.image(shots / "site-01-list.png", "图 2  现场管理列表：按名称、类型、区域和状态筛选真实页面")
    m.list([
        "进入现场融合管理 → 现场管理，使用名称、区域、类型或状态筛选。",
        "点击重置清除条件；导出前先限定范围并核对结果数量。",
        "点击现场名称或配置入口进入详情与融合配置。",
    ], numbered=True)
    m.h2("3.2 新增或编辑现场")
    m.image(shots / "site-02-add.png", "图 3  新增现场：分区填写基础资料、业务属性和备注")
    m.list([
        "点击新增，填写现场名称、编码、类型、区域和地址。",
        "联系人、电话等属于业务敏感信息，仅按授权录入；测试环境使用演示数据。",
        "保存后回到列表确认记录，再进入配置画布补充平台、组织和设备关系。",
    ], numbered=True)
    m.h2("3.3 批量导入现场")
    m.image(shots / "site-03-import.png", "图 4  现场批量导入：下载模板、上传 Excel 并查看校验结果")
    m.list([
        "先下载最新版模板，不自行更改列名、必填标识和枚举值。",
        "使用 3—5 条演示记录试导；校验失败时按提示修改原文件，不重复导入已成功记录。",
        "正式导入后按现场编码和名称抽查，并保存导入结果作为交付证据。",
    ], numbered=True)
    m.h2("3.4 融合配置画布")
    m.image(shots / "site-04-canvas.png", "图 5  现场融合配置画布：集中查看平台、组织与节点关系")
    m.p("画布按现场维度聚合主平台、子平台、组织、联系人、服务器和设备位置。先保存节点基础信息，再建立关系；调整关系前记录原拓扑，避免误断业务链路。")
    m.image(shots / "site-05-canvas-vertical.png", "图 6  纵向布局：适合查看平台到组织、服务器的上下游层级")
    m.image(shots / "site-06-canvas-fullscreen.png", "图 7  全屏布局：在复杂现场中缩放、定位并检查节点详情")
    m.list([
        "用横向/纵向布局切换适配现场规模；全屏模式用于复杂关系检查。",
        "点击节点查看右侧详情，核对名称、编码、状态、归属和关联数量。",
        "新增或解除关系后保存，再刷新页面确认后端数据已生效。",
        "发现重复节点时先确认引用关系，优先合并或停用，不直接删除。",
    ], numbered=True)
    m.h2("3.5 服务器与凭据")
    m.image(shots / "site-07-server.png", "图 8  服务器管理：维护地址、用途、状态、凭据和关联信息")
    m.list([
        "新增服务器时填写名称、IP、端口、用途、操作系统和所属现场。",
        "登录凭据由加密密钥保护；查看、修改或批量维护必须使用独立权限并留痕。",
        "设备机房、机柜和位置变更后，同步核对硬件资产关系。",
        "删除服务器前检查自动巡检目标和现场拓扑引用。",
    ], numbered=True)
    m.h2("3.6 批量导入服务器")
    m.image(shots / "site-08-server-import.png", "图 9  服务器批量导入：下载模板后上传并核对导入范围")
    m.list([
        "模板中的现场编码必须与系统已有现场一致。",
        "IP、端口和用途按现场资产台账填写；凭据列只允许授权人员处理。",
        "导入后检查重复地址、空凭据和未关联现场，并用单台服务器验证连接。",
    ], numbered=True)
    m.h2("3.7 功能版本记录")
    m.image(shots / "site-09-version.png", "图 10  功能版本记录：查看现场融合功能变更、版本和说明")
    m.p("版本记录用于确认现场看到的页面能力。排查“页面与手册不一致”时，先核对当前前端版本、后端 JAR 校验值和功能版本记录，再判断是权限、缓存还是版本差异。")

    m.h1("4. 现场自动巡检模块")
    m.p("现场自动巡检由工具、模板、目标、计划和记录组成。推荐顺序是：先维护现场服务器，再建工具和模板，选择目标，用单目标手工执行验证，最后创建停用计划并在维护窗口启用。")
    m.h2("4.1 总览与巡检记录")
    m.image(shots / "auto-01-overview.png", "图 11  自动巡检总览与记录：查看状态、成功率和最近任务")
    m.p("先通过总览判断成功率、失败目标和最近运行，再打开记录详情定位具体步骤。统计卡用于快速筛选，不替代逐条结果核验。")
    m.image(shots / "auto-02-dashboard.png", "图 12  巡检分析抽屉：从日期、成功率和趋势下钻到目标结果")
    m.h2("4.2 周报与月报导出")
    m.image(shots / "auto-03-week-report.png", "图 13  周报导出：选择时间范围并生成巡检统计文件")
    m.image(shots / "auto-04-month-report.png", "图 14  月报导出：按自然月汇总执行次数、成功率和失败项")
    m.list([
        "导出前确认时间范围、现场范围和时区；跨月数据分别导出再汇总。",
        "导出后抽查总次数、成功数、失败数是否与页面筛选结果一致。",
        "报告中的设备地址、输出日志按内部敏感资料保存，不通过公共渠道传输。",
    ], numbered=True)
    m.h2("4.3 记录详情")
    m.image(shots / "auto-05-record-detail.png", "图 15  巡检记录详情：按目标和步骤查看状态、耗时、输出与失败原因")
    m.list([
        "先看整次任务状态，再展开失败目标，最后定位具体失败步骤。",
        "区分网络不可达、认证失败、命令失败、输出解析失败和超时。",
        "保存任务 ID、目标 ID、步骤名、时间和脱敏日志；工单中不粘贴口令。",
        "修复后优先重跑单个目标，验证通过后再恢复批量计划。",
    ], numbered=True)
    m.h2("4.4 巡检配置入口与内置指引")
    m.image(shots / "auto-06-config.png", "图 16  巡检配置入口：集中切换模板、工具、目标和计划")
    m.image(shots / "auto-07-guide.png", "图 17  页面内操作指引：查看配置顺序、字段说明和注意事项")
    m.h2("4.5 模板与步骤")
    m.image(shots / "auto-08-template.png", "图 18  巡检模板编辑：组织步骤、顺序、参数、超时和失败策略")
    m.list([
        "进入巡检配置 → 模板管理，新建模板并填写名称、用途和说明。",
        "按执行顺序添加步骤，选择工具，填写超时、失败策略和参数。",
        "先保存为停用或草稿状态，用单个演示目标手工验证，再开放给计划。",
        "修改已被计划使用的模板前，复制新版本并保留旧版，避免影响正在运行的任务。",
    ], numbered=True)
    m.image(shots / "auto-09-tool-picker.png", "图 19  工具选择器：从受控工具库选择巡检能力")
    m.h2("4.6 HTTP 步骤配置")
    m.image(shots / "auto-10-http-step.png", "图 20  HTTP 步骤：配置请求方法、地址、请求头、参数和断言")
    m.list([
        "地址优先引用变量或目标字段，避免把生产地址写死在共享模板。",
        "请求头中的 Token、Cookie 等敏感值使用凭据变量，不写入步骤说明。",
        "设置合理超时与成功码；响应体断言只匹配稳定字段，避免整段文本硬匹配。",
    ], numbered=True)
    m.h2("4.7 目标与凭据")
    m.p("目标可关联服务器或独立地址。优先引用现场融合中的服务器资产，减少重复维护。用户名、密码、密钥等由后端加密保存；列表不得明文展示，导出文件不得包含凭据。")
    m.callout("凭据管理", "批量设置凭据前核对目标范围；密钥变更必须先备份并验证历史数据可解密。", "red")
    m.h2("4.8 计划创建与启停")
    m.image(shots / "auto-11-plan-list.png", "图 21  巡检计划列表：查看调度状态、模板、目标和最近执行")
    m.image(shots / "auto-12-plan-dialog.png", "图 22  新增巡检计划：选择模板、目标、调度表达式和启停状态")
    m.list([
        "点击新增计划，选择已验证的模板与目标集合。",
        "设置 Cron 或调度周期，确认时区、首次执行时间和并发影响。",
        "新计划默认停用；先手动执行一次并核对结果，再在维护窗口启用。",
        "停用计划不会继续调度，但已开始的任务可能仍会完成；必要时检查任务状态。",
    ], numbered=True)
    m.h2("4.9 服务步骤与目标范围")
    m.image(shots / "auto-13-service-step.png", "图 23  服务检查步骤：配置服务名称、检查方式、期望状态和失败策略")
    m.image(shots / "auto-14-target-tree.png", "图 24  现场—服务器目标树：按现场批量选择巡检目标")
    m.image(shots / "auto-15-site-relation.png", "图 25  现场服务器关系页：巡检目标与现场资产的来源页面")
    m.list([
        "优先从现场—服务器树选择目标，避免同一设备以多个地址重复纳管。",
        "批量选择前先限定现场，展开核对服务器数量、状态和用途。",
        "现场资产发生停用、地址或凭据变更时，同步检查巡检目标。",
        "服务检查需要现场确认服务名、运行用户和权限，先在单台服务器验证。",
    ], numbered=True)

    m.h1("5. 典型操作路径")
    m.table(["场景", "推荐路径", "完成标准"], [
        ["新现场上线", "新增现场 → 配置平台/组织 → 录入服务器 → 核对拓扑", "关系完整、无重复资产"],
        ["新增巡检", "建模板 → 配步骤 → 建目标 → 手工验证 → 建停用计划 → 启用", "单目标通过、窗口已确认"],
        ["服务器批量纳管", "下载模板 → 小批试导 → 正式导入 → 现场关系核对", "无重复、归属正确"],
        ["故障复盘", "记录详情 → 失败步骤 → 现场/服务器关系核查", "原因、修复、复测证据齐全"],
    ], [1800, 4300, 3260])

    m.h1("6. 权限与数据安全")
    m.list([
        "生产账号一人一号，不共享管理员账号；离岗立即停用。",
        "凭据查看、批量修改、删除、导入、导出和巡检启停设置为高风险权限。",
        "演示数据与真实数据使用不同命名和网段；测试环境禁止复制生产口令。",
        "导出的现场资料、联系方式和巡检日志按敏感资料保存，设置有效期和访问范围。",
        "定期审查登录日志、操作日志、失败任务和长期未使用账号。",
    ])

    m.h1("7. 常见问题")
    m.table(["问题", "检查", "建议"], [
        ["看不到菜单/按钮", "角色、菜单权限、重新登录", "由管理员按岗位授权"],
        ["现场保存失败", "必填项、编码唯一性、关联对象", "保留提示与网络请求编号"],
        ["现场关系显示不全", "节点状态、关联关系、画布筛选", "刷新后逐级核对平台、组织和服务器"],
        ["服务器导入失败", "模板版本、必填列、现场编码", "先用少量记录试导并保存错误明细"],
        ["巡检认证失败", "目标地址、账号、密钥、授权范围", "使用单目标验证，禁止反复锁定账号"],
        ["计划未触发", "状态、Cron、时区、Quartz 日志", "先手工执行，再检查调度"],
    ], [2100, 3300, 3960])

    m.h1("8. 操作验收记录模板")
    m.table(["字段", "填写内容"], [
        ["操作日期/人员", ""], ["环境/版本/SHA-256", ""], ["模块与功能", ""],
        ["对象范围", ""], ["操作前状态", ""], ["操作结果", ""],
        ["截图/日志位置", ""], ["异常与回滚", ""], ["复核人", ""],
    ], [2800, 6560])
    m.save(out)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--screenshots", type=Path, required=True)
    parser.add_argument("--deployment-screenshots", type=Path, required=True)
    args = parser.parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    build_deployment(
        args.output_dir / "华东信息融合平台-前后端程序部署手册-v3.9.1.docx",
        args.deployment_screenshots,
    )
    build_operation(args.output_dir / "华东信息融合平台-操作手册-v3.9.1.docx", args.screenshots)
    print("documents_built=2")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
